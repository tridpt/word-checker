"""Helper dung chung cho cac test."""

import os
import sys

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Cho phep import package 'checker' va 'config' tu thu muc goc du an
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))


def add_para(doc, text, font="Times New Roman", size=14,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5):
    """Them mot doan voi dinh dang cho truoc."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    if font is not None:
        r.font.name = font
    if size is not None:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if spacing is not None:
        p.paragraph_format.line_spacing = spacing
    return p


@pytest.fixture
def make_doc(tmp_path):
    """Tra ve ham build(paragraphs) -> duong dan file .docx tam."""
    def build(paragraphs):
        doc = Document()
        for item in paragraphs:
            if isinstance(item, str):
                add_para(doc, item)
            else:
                add_para(doc, **item)
        out = os.path.join(tmp_path, "test.docx")
        doc.save(out)
        return out
    return build
