"""Test soat loi trong bang, dau/chan trang, text box (ngoai than bai chinh)."""

import os

from docx import Document

import config
from checker.document import load_document
from checker.runner import run_checks
from checker.spelling import check_spelling
from checker.stats import compute_stats
from checker.textcheck import check_text


def _make_doc_with_table(tmp_path, cell_texts):
    """Tao .docx co 1 doan body + 1 bang chua cac o cell_texts."""
    doc = Document()
    doc.add_paragraph("Doan than bai binh thuong.")
    table = doc.add_table(rows=1, cols=len(cell_texts))
    for i, text in enumerate(cell_texts):
        table.rows[0].cells[i].text = text
    out = os.path.join(tmp_path, "with_table.docx")
    doc.save(out)
    return out


def _make_doc_with_header_footer(tmp_path, header_text, footer_text):
    doc = Document()
    doc.add_paragraph("Doan than bai.")
    section = doc.sections[0]
    section.header.paragraphs[0].text = header_text
    section.footer.paragraphs[0].text = footer_text
    out = os.path.join(tmp_path, "with_header.docx")
    doc.save(out)
    return out


def test_extra_segments_collected_from_table(tmp_path):
    path = _make_doc_with_table(tmp_path, ["Ho ten", "Dia chi"])
    doc = load_document(path)
    # body chi co 1 doan, cac o bang nam trong extra_segments
    assert len(doc.paragraphs) >= 1
    seg_texts = [s.text for s in doc.extra_segments]
    assert "Ho ten" in seg_texts
    assert "Dia chi" in seg_texts
    assert all(s.location == "Bang" for s in doc.extra_segments)


def test_spelling_detected_in_table(tmp_path):
    path = _make_doc_with_table(tmp_path, ["Can bổ xung thong tin"])
    doc = load_document(path)
    issues = check_spelling(doc)
    assert any("bổ xung" in it.message or "Bang" in it.message for it in issues)
    # loi trong bang la loi cap tai lieu (paragraph None) va co nhan vi tri
    table_issues = [it for it in issues if it.paragraph is None and "Bang" in it.message]
    assert table_issues


def test_text_mechanics_detected_in_table(tmp_path):
    path = _make_doc_with_table(tmp_path, ["Noi dung  co cach doi , sai"])
    doc = load_document(path)
    issues = check_text(doc, config.TEXT_CHECKS, config.MAX_CONSECUTIVE_EMPTY)
    table_issues = [it for it in issues if "Bang" in it.message]
    assert any("cach doi" in it.message for it in table_issues)


def test_spelling_detected_in_header_footer(tmp_path):
    path = _make_doc_with_header_footer(tmp_path, "Bao cao nghành X", "Ghi chu chia sẽ")
    doc = load_document(path)
    issues = check_spelling(doc)
    msgs = " ".join(it.message for it in issues)
    assert "Dau trang" in msgs
    assert "Chan trang" in msgs


def test_stats_ignore_extra_segments(tmp_path):
    """Thong ke chi tinh than bai, khong tinh o bang -> so tu on dinh."""
    path = _make_doc_with_table(tmp_path, ["Mot hai ba bon nam sau bay tam chin muoi"])
    stats = compute_stats(load_document(path))
    # body chi co 'Doan than bai binh thuong.' = 4 tu that (chuoi o bang bi bo qua)
    assert stats["words"] == len("Doan than bai binh thuong.".split())


def test_clean_table_has_no_issues(tmp_path):
    path = _make_doc_with_table(tmp_path, ["Ho ten", "Dia chi", "So dien thoai"])
    profile = config.get_profile(None)
    issues = run_checks(path, profile)
    # cac o bang sach se khong sinh loi van ban/chinh ta
    extra_issues = [it for it in issues if "Bang" in it.message]
    assert not extra_issues
