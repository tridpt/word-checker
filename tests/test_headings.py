"""Test kiem tra tieu de muc va danh so de muc."""

import os

from docx import Document

from checker.document import load_document
from checker.headings import check_headings


def _build(tmp_path, items):
    """items: list cua (text, level). level 0 = doan thuong."""
    doc = Document()
    for text, level in items:
        if level == 0:
            doc.add_paragraph(text)
        else:
            doc.add_heading(text, level=level)
    out = os.path.join(tmp_path, "h.docx")
    doc.save(out)
    return out


def _messages(path):
    return [it.message for it in check_headings(load_document(path), None)]


def test_level_jump_detected(tmp_path):
    path = _build(tmp_path, [
        ("Chuong 1", 1),
        ("Muc con sau", 3),  # nhay tu 1 sang 3
    ])
    assert any("Nhay cap tieu de" in m for m in _messages(path))


def test_no_level_jump_when_sequential(tmp_path):
    path = _build(tmp_path, [
        ("Chuong 1", 1),
        ("Muc 1.1", 2),
        ("Muc con", 3),
    ])
    assert not any("Nhay cap tieu de" in m for m in _messages(path))


def test_numbering_gap_detected(tmp_path):
    path = _build(tmp_path, [
        ("1. Mo dau", 1),
        ("2. Noi dung", 1),
        ("4. Ket luan", 1),  # thieu 3
    ])
    msgs = _messages(path)
    assert any("khong lien tuc" in m and "'4'" in m for m in msgs)


def test_numbering_ok_when_sequential(tmp_path):
    path = _build(tmp_path, [
        ("1. Mo dau", 1),
        ("2. Noi dung", 1),
        ("3. Ket luan", 1),
    ])
    assert not any("khong lien tuc" in m for m in _messages(path))


def test_subsection_numbering(tmp_path):
    path = _build(tmp_path, [
        ("1. Tong quan", 1),
        ("1.1 Boi canh", 2),
        ("1.3 Muc tieu", 2),  # thieu 1.2
    ])
    assert any("khong lien tuc" in m for m in _messages(path))
