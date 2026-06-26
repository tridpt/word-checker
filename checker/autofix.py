"""Tu dong sua mot so loi co hoc va loi chinh ta, luu ra file .docx moi.

Nguyen tac an toan:
- Khong ghi de file goc; luu ra file moi (mac dinh them hau to '_fixed').
- Sua o cap 'run' de giu nguyen dinh dang (font, mau, dam/nghieng...).
- Chi sua nhung loi chac chan (whitespace, dau cau, loi chinh ta trong tu dien).
"""

import re
from collections import Counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .spelling import load_typo_dict

_PUNCT_NO_SPACE_BEFORE = ",.;:!?)"

_ALIGN_TO_ENUM = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _is_heading(paragraph) -> bool:
    name = paragraph.style.name if paragraph.style else ""
    return name.startswith("Heading") or name in ("Title", "Subtitle")


def _fix_run_text(text: str, typos: dict[str, str]) -> tuple[str, int]:
    """Sua loi co hoc + chinh ta trong text cua mot run. Tra ve (text_moi, so_loi_sua)."""
    count = 0

    # 1) Bo dau cach truoc dau cau
    new, n = re.subn(r"[ \t]+([" + re.escape(_PUNCT_NO_SPACE_BEFORE) + r"])", r"\1", text)
    count += n
    text = new

    # 2) Gop nhieu dau cach lien tiep thanh mot
    new, n = re.subn(r"  +", " ", text)
    count += n
    text = new

    # 3) Them dau cach sau , ; : neu thieu (tru giua hai chu so)
    def _add_space(m):
        return m.group(1) + " "

    new, n = re.subn(r"([,;:])(?=[^\s\d])", _add_space, text)
    # tranh dong cham truong hop so:so (gio) -> dieu kien tren da loai chu so phia sau
    count += n
    text = new

    # 4) Sua loi chinh ta theo tu dien (khop ca tu, khong phan biet hoa/thuong)
    for wrong, right in typos.items():
        pattern = re.compile(r"(?<!\w)(" + re.escape(wrong) + r")(?!\w)", re.IGNORECASE | re.UNICODE)

        def _repl(m):
            nonlocal count
            count += 1
            return _match_case(m.group(1), right)

        text = pattern.sub(_repl, text)

    return text, count


def _match_case(found: str, replacement: str) -> str:
    """Giu kieu viet hoa: neu tu sai viet hoa chu cai dau thi tu dung cung vay."""
    if found[:1].isupper():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def _remove_extra_empty_paragraphs(document, max_consecutive: int) -> int:
    """Xoa cac doan trong lien tiep vuot qua gioi han. Tra ve so doan da xoa."""
    removed = 0
    run_len = 0
    to_delete = []
    for p in document.paragraphs:
        if not p.text.strip():
            run_len += 1
            if run_len > max_consecutive:
                to_delete.append(p)
        else:
            run_len = 0
    for p in to_delete:
        el = p._element
        el.getparent().remove(el)
        removed += 1
    return removed


def _fix_formatting(document, profile: dict) -> int:
    """Ap dat dinh dang theo profile cho cac doan than bai. Tra ve so thay doi."""
    changes = 0
    allowed_fonts = list(profile.get("body_font_names") or [])
    allowed_sizes = [float(s) for s in (profile.get("body_font_sizes") or [])]
    align = profile.get("body_alignment")
    spacing = profile.get("line_spacing")
    fli = profile.get("first_line_indent_cm")
    sb = profile.get("space_before_pt")
    sa = profile.get("space_after_pt")

    body = [p for p in document.paragraphs if p.text.strip() and not _is_heading(p)]

    # Chon font/co chu DICH = gia tri cho phep dang duoc dung nhieu nhat trong body
    # (tranh dat ve gia tri dau danh sach gay lech voi phan con lai).
    font_used: Counter = Counter()
    size_used: Counter = Counter()
    for p in body:
        for run in p.runs:
            if not run.text.strip():
                continue
            if run.font.name in allowed_fonts:
                font_used[run.font.name] += 1
            if run.font.size and run.font.size.pt in allowed_sizes:
                size_used[run.font.size.pt] += 1

    target_font = (font_used.most_common(1)[0][0] if font_used
                   else (allowed_fonts[0] if allowed_fonts else None))
    target_size = (size_used.most_common(1)[0][0] if size_used
                   else (allowed_sizes[0] if allowed_sizes else None))
    allowed_fonts_set = set(allowed_fonts)
    allowed_sizes_set = set(allowed_sizes)

    for p in body:
        for run in p.runs:
            if not run.text.strip():
                continue
            if target_font and (run.font.name is None or run.font.name not in allowed_fonts_set):
                if run.font.name != target_font:
                    run.font.name = target_font
                    changes += 1
            if target_size is not None:
                cur = run.font.size.pt if run.font.size else None
                if cur is None or cur not in allowed_sizes_set:
                    run.font.size = Pt(target_size)
                    changes += 1

        pf = p.paragraph_format
        if align in _ALIGN_TO_ENUM and p.alignment != _ALIGN_TO_ENUM[align]:
            p.alignment = _ALIGN_TO_ENUM[align]
            changes += 1
        if spacing is not None and pf.line_spacing != spacing:
            pf.line_spacing = spacing
            changes += 1
        if fli is not None:
            pf.first_line_indent = Cm(fli)
        if sb is not None:
            pf.space_before = Pt(sb)
        if sa is not None:
            pf.space_after = Pt(sa)

    changes += _fix_margins(document, profile)
    return changes


def _fix_margins(document, profile: dict) -> int:
    margins = profile.get("margins_cm")
    if not margins:
        return 0
    tol = profile.get("margin_tolerance_cm", 0.2)
    changes = 0
    try:
        section = document.sections[0]
    except (IndexError, AttributeError):
        return 0
    mapping = {
        "top": "top_margin",
        "bottom": "bottom_margin",
        "left": "left_margin",
        "right": "right_margin",
    }
    for side, attr in mapping.items():
        req = margins.get(side)
        if req is None:
            continue
        cur = getattr(section, attr)
        cur_cm = cur.cm if cur is not None else None
        if cur_cm is None or abs(cur_cm - req) > tol:
            setattr(section, attr, Cm(req))
            changes += 1
    return changes


def autofix(
    src_path: str,
    out_path: str,
    max_consecutive_empty: int = 1,
    fix_spelling: bool = True,
    fix_format: bool = False,
    profile: dict | None = None,
) -> dict:
    """Sua file va luu ra out_path. Tra ve thong ke so loi da sua."""
    document = Document(src_path)
    typos = load_typo_dict() if fix_spelling else {}

    # Sua cac doan cuoi cau bi thua khoang trang: xu ly o run cuoi cung
    text_fixes = 0
    for p in document.paragraphs:
        runs = p.runs
        for i, run in enumerate(runs):
            new_text, n = _fix_run_text(run.text, typos)
            # Cat khoang trang cuoi doan: chi ap dung cho run cuoi cung co noi dung
            if i == len(runs) - 1:
                new_text = new_text.rstrip()
            if new_text != run.text:
                run.text = new_text
            text_fixes += n

    empty_removed = _remove_extra_empty_paragraphs(document, max_consecutive_empty)

    format_fixes = 0
    if fix_format and profile:
        format_fixes = _fix_formatting(document, profile)

    document.save(out_path)
    return {
        "text_and_spelling_fixes": text_fixes,
        "empty_paragraphs_removed": empty_removed,
        "format_fixes": format_fixes,
        "out_path": out_path,
    }
