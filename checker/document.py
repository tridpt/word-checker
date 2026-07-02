"""Doc file .docx va trich xuat thong tin tung doan van (paragraph).

Xu ly viec phan giai font/co chu "hieu luc" (effective): mot run co the lay
font tu chinh no, tu style cua doan, tu style goc (based_on), hoac tu mac dinh
cua tai lieu (docDefaults).
"""

from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


@dataclass
class RunInfo:
    text: str
    font_name: str | None
    font_size: float | None  # point


@dataclass
class ParagraphInfo:
    index: int           # so thu tu doan trong tai lieu
    number: int          # so thu tu doan KHONG rong (de hien thi cho nguoi dung)
    text: str
    style_name: str
    is_heading: bool
    alignment: str | None       # "left" | "center" | "right" | "justify" | None
    line_spacing: float | None  # vd 1.5; None neu khong xac dinh
    first_line_indent_cm: float | None = None  # thut le dong dau (cm)
    space_before_pt: float | None = None       # khoang cach truoc doan (pt)
    space_after_pt: float | None = None         # khoang cach sau doan (pt)
    fonts: set[str] = field(default_factory=set)        # cac font hieu luc trong doan
    sizes: set[float] = field(default_factory=set)      # cac co chu hieu luc trong doan
    runs: list[RunInfo] = field(default_factory=list)


@dataclass
class ExtraSegment:
    """Mot doan van ban nam ngoai than bai chinh (bang, dau/chan trang, text box).

    Chi kiem tra loi co hoc va chinh ta tren cac doan nay; KHONG kiem tra dinh
    dang (font/co chu) vi bang bieu va dau/chan trang thuong dung dinh dang rieng
    mot cach hop le -> tranh bao nham.
    """
    text: str
    location: str  # nhan vi tri, vd "Bang", "Dau trang", "Chan trang", "Text box"


@dataclass
class DocInfo:
    path: str
    paragraphs: list[ParagraphInfo]
    margins_cm: dict           # {"top","bottom","left","right"} hoac {} neu khong doc duoc
    default_font: str | None
    default_size: float | None
    extra_segments: list[ExtraSegment] = field(default_factory=list)


_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def _doc_default_font(document) -> tuple[str | None, float | None]:
    """Doc font/co mac dinh cua tai lieu tu w:docDefaults."""
    name, size = None, None
    try:
        styles_el = document.styles.element
        rpr = styles_el.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                name = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
            sz = rpr.find(qn("w:sz"))
            if sz is not None and sz.get(qn("w:val")):
                size = int(sz.get(qn("w:val"))) / 2.0  # half-points -> points
    except Exception:
        pass
    return name, size


def _style_font(style) -> tuple[str | None, float | None]:
    """Lay font/co chu tu chuoi style (theo based_on), tra ve gia tri dau tien tim thay."""
    name, size = None, None
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        try:
            if name is None and style.font.name:
                name = style.font.name
            if size is None and style.font.size:
                size = style.font.size.pt
        except Exception:
            pass
        if name and size:
            break
        style = getattr(style, "base_style", None)
    return name, size


def _resolve_run(run, paragraph, defaults) -> RunInfo:
    default_name, default_size = defaults
    name = run.font.name
    size = run.font.size.pt if run.font.size else None

    if name is None or size is None:
        # thu lay tu style cua run (character style)
        try:
            rs_name, rs_size = _style_font(run.style)
        except Exception:
            rs_name, rs_size = None, None
        name = name or rs_name
        size = size or rs_size

    if name is None or size is None:
        ps_name, ps_size = _style_font(paragraph.style)
        name = name or ps_name
        size = size or ps_size

    name = name or default_name
    size = size or default_size
    return RunInfo(text=run.text, font_name=name, font_size=size)


def _line_spacing(paragraph) -> float | None:
    pf = paragraph.paragraph_format
    if pf.line_spacing is None:
        return None
    val = pf.line_spacing
    # Neu la multiple (1.0, 1.5...) python-docx tra ve float.
    # Neu la khoang cach tuyet doi (Pt) thi tra ve Length -> chuyen ve None de bo qua.
    if isinstance(val, float):
        return round(val, 3)
    return None


def _para_indent_spacing(paragraph):
    """Tra ve (first_line_indent_cm, space_before_pt, space_after_pt)."""
    pf = paragraph.paragraph_format
    fli = pf.first_line_indent.cm if pf.first_line_indent is not None else None
    sb = pf.space_before.pt if pf.space_before is not None else None
    sa = pf.space_after.pt if pf.space_after is not None else None
    return (
        round(fli, 2) if fli is not None else None,
        round(sb, 2) if sb is not None else None,
        round(sa, 2) if sa is not None else None,
    )


def _iter_table_texts(table):
    """Duyet de quy moi o trong bang (ke ca bang long nhau), tra ve text tung doan."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    yield p.text
            for nested in cell.tables:
                yield from _iter_table_texts(nested)


def _collect_extra_segments(document) -> list["ExtraSegment"]:
    """Thu thap text ngoai than bai: bang, dau/chan trang, text box."""
    segments: list[ExtraSegment] = []

    # 1) Bang trong than tai lieu
    try:
        for table in document.tables:
            for text in _iter_table_texts(table):
                segments.append(ExtraSegment(text=text, location="Bang"))
    except Exception:
        pass

    # 2) Dau trang / chan trang cua tat ca cac section
    try:
        for section in document.sections:
            for header, label in (
                (section.header, "Dau trang"),
                (section.first_page_header, "Dau trang (trang dau)"),
                (section.even_page_header, "Dau trang (trang chan)"),
                (section.footer, "Chan trang"),
                (section.first_page_footer, "Chan trang (trang dau)"),
                (section.even_page_footer, "Chan trang (trang chan)"),
            ):
                if header is None or getattr(header, "is_linked_to_previous", False):
                    continue
                for p in header.paragraphs:
                    if p.text.strip():
                        segments.append(ExtraSegment(text=p.text, location=label))
                for table in header.tables:
                    for text in _iter_table_texts(table):
                        segments.append(ExtraSegment(text=text, location=label))
    except Exception:
        pass

    # 3) Text box / shape: quet text trong cac phan tu w:txbxContent tren toan XML
    try:
        seen_txbx = set()
        for txbx in document.element.iter(qn("w:txbxContent")):
            texts = []
            for t in txbx.iter(qn("w:t")):
                if t.text:
                    texts.append(t.text)
            joined = "".join(texts).strip()
            if joined and joined not in seen_txbx:
                seen_txbx.add(joined)
                segments.append(ExtraSegment(text=joined, location="Text box"))
    except Exception:
        pass

    return segments


def load_document(path: str) -> DocInfo:
    document = Document(path)
    defaults = _doc_default_font(document)

    paragraphs: list[ParagraphInfo] = []
    number = 0
    for i, p in enumerate(document.paragraphs):
        style_name = p.style.name if p.style else ""
        is_heading = style_name.startswith("Heading") or style_name in ("Title", "Subtitle")

        align = _ALIGN_MAP.get(p.alignment) if p.alignment is not None else None

        runs = [_resolve_run(r, p, defaults) for r in p.runs]
        fonts = {r.font_name for r in runs if r.text.strip() and r.font_name}
        sizes = {r.font_size for r in runs if r.text.strip() and r.font_size}

        text = p.text
        if text.strip():
            number += 1

        fli, sb, sa = _para_indent_spacing(p)
        paragraphs.append(
            ParagraphInfo(
                index=i,
                number=number,
                text=text,
                style_name=style_name,
                is_heading=is_heading,
                alignment=align,
                line_spacing=_line_spacing(p),
                first_line_indent_cm=fli,
                space_before_pt=sb,
                space_after_pt=sa,
                fonts=fonts,
                sizes=sizes,
                runs=runs,
            )
        )

    margins_cm = {}
    try:
        section = document.sections[0]
        margins_cm = {
            "top": round(section.top_margin.cm, 2),
            "bottom": round(section.bottom_margin.cm, 2),
            "left": round(section.left_margin.cm, 2),
            "right": round(section.right_margin.cm, 2),
        }
    except Exception:
        margins_cm = {}

    return DocInfo(
        path=path,
        paragraphs=paragraphs,
        margins_cm=margins_cm,
        default_font=defaults[0],
        default_size=defaults[1],
        extra_segments=_collect_extra_segments(document),
    )
