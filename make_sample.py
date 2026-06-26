"""Tao mot file Word mau co san loi de thu nghiem tool."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build(path: str = "sample.docx"):
    doc = Document()

    # Tieu de (heading) - khong bi kiem tra body
    h = doc.add_heading("Bao cao thuc tap", level=1)

    # Doan dung chuan: Times New Roman 14, canh deu, gian dong 1.5
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Day la doan van mau dung chuan dinh dang.")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.line_spacing = 1.5

    # Doan sai font (Arial) va sai co chu (11)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Doan nay dung sai font va sai co chu so voi yeu cau.")
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.line_spacing = 1.0

    # Doan co loi co hoc van ban: cach doi, cach truoc dau phay, thieu cach sau dau cham phay
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Cau nay co  cach doi , va thieu cach sau dau cham phay;tiep theo nhu the.   ")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(14)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.line_spacing = 1.5

    # Doan co loi chinh ta tu tu dien
    p4 = doc.add_paragraph()
    r4 = p4.add_run("Toi muon chia sẽ kinh nghiem va bổ xung kien thuc cho moi nguời.")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(14)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.paragraph_format.line_spacing = 1.5

    # Hai doan trong lien tiep
    doc.add_paragraph("")
    doc.add_paragraph("")

    p5 = doc.add_paragraph()
    r5 = p5.add_run("Doan ket thuc bao cao.")
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(14)
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p5.paragraph_format.line_spacing = 1.5

    doc.save(path)
    print(f"Da tao file mau: {path}")


if __name__ == "__main__":
    build()
